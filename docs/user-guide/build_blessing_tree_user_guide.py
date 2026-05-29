from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_blessing_tree_user_guide_pdf import FIELD_REFERENCE


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "user-guide"
SCREENSHOT_DIR = OUT_DIR / "screenshots"
DOCX_PATH = OUT_DIR / "Blessing Tree User Guide.docx"

BRAND_PURPLE = "2A0F3D"
BRAND_GOLD = "D8A928"
BRAND_GREEN = "2C7A52"
SOFT_FILL = "F8F5EF"
TABLE_HEADER = "EDE6F3"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold


def set_cell_width(cell, width_inches: float) -> None:
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_horizontal_rule(paragraph, color: str = BRAND_GOLD) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(BRAND_PURPLE if level <= 2 else "1F4D78")
    return paragraph


def add_body(doc: Document, text: str):
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    return paragraph


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.add_run(item)


def add_steps(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.add_run(item)


def add_note(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.allow_autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, SOFT_FILL)
    set_cell_width(cell, 6.25)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(BRAND_PURPLE)
    paragraph = cell.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    table.allow_autofit = False
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, TABLE_HEADER)
        if widths:
            set_cell_width(cell, widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            if widths:
                set_cell_width(cells[idx], widths[idx])
    doc.add_paragraph()


def add_screenshot(doc: Document, filename: str, caption: str, width: float = 6.25) -> None:
    path = SCREENSHOT_DIR / filename
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    caption_paragraph = doc.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(8)
    for run in caption_paragraph.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(90, 90, 90)


def add_workflow(doc: Document, title: str, steps: list[str]) -> None:
    add_heading(doc, title, level=3)
    add_steps(doc, steps)


def add_field_reference_section(doc: Document) -> None:
    add_heading(doc, "Detailed Field Reference", level=1)
    add_body(
        doc,
        "This section explains what each major screen field is for, how staff should use it, and practical suggestions for clean data entry. Treat this as the operator reference when training staff or troubleshooting confusing screens.",
    )
    for title, intro, rows in FIELD_REFERENCE:
        add_heading(doc, title, level=2)
        add_body(doc, intro)
        add_table(doc, ["Field", "What it does", "Suggestion"], rows, widths=[1.25, 2.25, 2.75])


def build_document() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[style_name].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor.from_string(BRAND_PURPLE)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor.from_string(BRAND_PURPLE)
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.color.rgb = RGBColor.from_string("1F4D78")

    logo = ROOT / "blessing-tree-ui" / "public" / "blessing-tree-logo.png"
    if logo.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(logo), width=Inches(2.1))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Blessing Tree User Guide")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor.from_string(BRAND_PURPLE)
    add_horizontal_rule(title)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Campaign setup, people intake, sponsors, gifts, reporting, and administration")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(80, 80, 80)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Prepared {date.today().strftime('%B %-d, %Y')}")

    add_note(
        doc,
        "How to use this guide",
        "Start at Campaigns and work through the app from left to right. Screens and buttons may be hidden if your account does not have permission for that area.",
    )

    add_heading(doc, "Contents", level=1)
    add_bullets(
        doc,
        [
            "Getting started and navigation",
            "Campaigns and Campaign Studio",
            "People intake, directory, and reports",
            "Sponsors intake, directory, reports, and communications",
            "Gift search, operations, pool, status, QR scanning, and tag builder",
            "Ask Blessing Tree",
            "Administration, access, organization types, readiness rules, and health",
            "Common workflows and glossary",
        ],
    )

    doc.add_page_break()

    add_heading(doc, "Getting Started", level=1)
    add_body(
        doc,
        "Blessing Tree is organized around campaigns. A campaign is the operating container for a season or giving effort: its dates, teams, people served, sponsors, gifts, communications, reports, and readiness checks all belong to that campaign.",
    )
    add_screenshot(doc, "dashboard.png", "Dashboard with campaign switcher, widgets, Ask shortcuts, and readiness notes.")
    add_heading(doc, "Main Navigation", level=2)
    add_table(
        doc,
        ["Area", "What it is for"],
        [
            ["Dashboard", "Quick campaign snapshot, popular gift widgets, unsponsored gift counts, and pick-up-where-I-left-off prompts."],
            ["Ask Blessing Tree", "Plain-language help and reporting. Use it to ask where something is or to run campaign data questions."],
            ["Campaigns", "Campaign library, campaign overview, Campaign Studio, and Flyer Builder."],
            ["People", "Family, organization, child, adult, wishlist, and people reports."],
            ["Sponsors", "Sponsor intake, sponsor directory, sponsor reports, interaction logs, and sponsor email sends."],
            ["Gifts", "Gift search, reservations, operations, gift pool, Gift Status report, and Gift Tag Builder."],
            ["Admin", "User access, organization types, global campaign operation rules, LLM configuration, health, and app capabilities."],
        ],
        widths=[1.5, 4.75],
    )
    add_note(
        doc,
        "Campaign switcher",
        "The campaign dropdown in the top bar controls which campaign most screens use. If a report or workspace looks empty, first confirm the correct campaign is selected.",
    )

    add_heading(doc, "Campaigns", level=1)
    add_body(
        doc,
        "Campaigns define the season or effort being managed. The campaign library is where admins browse existing campaigns, open a campaign, or create a new one. A new campaign can be created from scratch or cloned from a previous campaign so teams, templates, milestones, and schedules do not have to be rebuilt every year.",
    )
    add_screenshot(doc, "campaign-library.png", "Campaign Library shows accessible campaigns and entry points into each campaign.")
    add_workflow(
        doc,
        "Create or Open a Campaign",
        [
            "Open Campaigns from the left navigation.",
            "Select an existing campaign card to open it, or choose the create action if you are an app admin.",
            "Use the campaign detail page for high-level settings, or open Campaign Studio for full setup.",
            "Set the campaign name, year, dates, lifecycle status, public sponsor slug, and campaign purpose.",
        ],
    )

    add_heading(doc, "Campaign Studio", level=1)
    add_body(
        doc,
        "Campaign Studio is the manager workspace for building and operating the campaign. It is divided into sections so the manager can move through setup without hunting across the app.",
    )
    add_screenshot(doc, "campaign-studio.png", "Campaign Studio overview with section rail, campaign summary, communications, schedule, and readiness.")
    add_table(
        doc,
        ["Studio Section", "Use it for"],
        [
            ["Overview", "Campaign snapshot, current counts, team summary, saved templates, upcoming schedule items, and readiness findings."],
            ["Team", "Campaign roster, campaign managers, coordinators, volunteers, teams, team roles, and app access for this campaign."],
            ["Communications", "Email templates, test emails, real sends, scheduled sends, merge fields, and communication history."],
            ["Schedule", "Calendar view for milestones, manual events, and scheduled communications."],
            ["Gift Rules", "Campaign-level rules such as sponsor gift limits, wishlist limits, and how many gifts define a fulfilled person."],
            ["Readiness", "Blockers, warnings, and notes that tell staff what must be fixed before a campaign is ready."],
            ["Settings", "Campaign metadata, lifecycle status, public sponsor settings, dates, campaign purpose, and flyer access."],
        ],
        widths=[1.55, 4.7],
    )
    add_heading(doc, "Teams and App Access", level=2)
    add_body(
        doc,
        "Teams are operational groups inside a campaign. A person can be on the campaign roster, can belong to one or more teams, and can have app access roles that control what screens they can see. Team roles describe what a person does operationally; app access controls what they can do in the software.",
    )
    add_bullets(
        doc,
        [
            "Campaign Manager can administer the campaign setup.",
            "People access is for intake, directory, and people reports.",
            "Sponsors access is for sponsor intake, directory, and sponsor reporting.",
            "Gift roles control gift search, operations, pool, wrapping, check-in, pickup, and distribution work.",
            "Reports access allows report viewing and exports without operational editing.",
        ],
    )
    add_heading(doc, "Milestones and Readiness", level=2)
    add_body(
        doc,
        "Milestones are dated campaign events such as registration opening, sponsor recruitment starting, gift intake ending, and pickup windows. Readiness rules use those milestones and campaign settings to warn staff about missing setup. Some readiness rules are blockers, meaning the campaign should not advance until they are fixed.",
    )
    add_bullets(
        doc,
        [
            "Milestone dates are managed in Campaign Studio Schedule and Readiness.",
            "Admin Campaign Operations controls the global milestone catalog and readiness rule definitions.",
            "Examples of blockers include missing sponsor recruitment dates, missing sponsor end dates, missing gift turn-in timing, and missing required campaign setup.",
            "Readiness findings include an action label so staff know which section to open.",
        ],
    )
    add_heading(doc, "Communications", level=2)
    add_body(
        doc,
        "Communications are campaign-scoped email templates. Templates can be used for manual sends, sponsor reminders, scheduled sends, and test emails. The template audience describes the intended recipient group; actual sends can be immediate, scheduled from a milestone, or locked to a sponsor when sent from that sponsor's drawer.",
    )
    add_bullets(
        doc,
        [
            "Use the template builder to create subject and body content.",
            "Use merge fields for campaign, sponsor, gift, recipient, and schedule values.",
            "Send a test email before sending a real campaign communication.",
            "Real sends are logged and can show recipient-level delivery status.",
            "Sponsor-specific sends can include gift lists such as gifts awaiting turn-in or received gifts.",
        ],
    )
    add_heading(doc, "Flyer Builder", level=2)
    add_body(
        doc,
        "The Flyer Builder creates sponsor-facing flyers for the campaign. It supports editable text and image placement, preview, and PDF-style output. Flyer QR codes should point sponsors to the public sponsor registration page for the selected campaign.",
    )

    add_heading(doc, "People", level=1)
    add_body(
        doc,
        "People is where staff enter and maintain the families, organizations, children, adults, contacts, pickup contacts, and wishlists served by the campaign.",
    )
    add_screenshot(doc, "people-intake.png", "People Intake starts new family, organization, child, adult, and wishlist workflows.")
    add_heading(doc, "People Intake", level=2)
    add_bullets(
        doc,
        [
            "Use Family/Household intake for families with children.",
            "Use Organization intake for nursing homes, foster care partners, mental health clients, homebound programs, or other partner groups.",
            "Organization Type controls whether the organization serves children, adults, or families.",
            "Children can use automatic labels such as Child One and Child Two so real names are not required.",
            "After adding a child or adult, open that person to add wishlist gifts, then return to add another person if needed.",
        ],
    )
    add_screenshot(doc, "people-directory.png", "People Directory is the maintenance screen for existing families, organizations, recipients, contacts, and wishlists.")
    add_heading(doc, "People Directory", level=2)
    add_body(
        doc,
        "Use the directory when someone is already in the system. Search for a family, organization, child, or adult, then click the row to open the drawer. The drawer is also where staff can add another child or adult to an existing household, add a family under an organization, edit contacts, and manage wishlist details.",
    )
    add_workflow(
        doc,
        "Add a Third Child to an Existing Household",
        [
            "Open People, then Directory.",
            "Search for the household or family.",
            "Click anywhere on the household row to open the drawer.",
            "Use Quick Actions or the Children section to choose Add Child.",
            "Save the child record, then open the child to add wishlist gifts.",
        ],
    )
    add_heading(doc, "People Reports", level=2)
    add_body(
        doc,
        "People Reports summarize group mix, people mix, wishlist readiness, gift workflow coverage, people still needing gifts, and pickup coordination. Reports include export buttons for Excel and PDF where available.",
    )

    add_heading(doc, "Sponsors", level=1)
    add_body(
        doc,
        "Sponsors are the people or organizations who commit to buying or donating gifts. Sponsor work includes intake, follow-up, gift commitments, public registration review, and sponsor-specific communications.",
    )
    add_screenshot(doc, "sponsor-directory.png", "Sponsors Directory shows sponsor records, contact status, gift commitments, and follow-up context.")
    add_heading(doc, "Sponsor Intake", level=2)
    add_bullets(
        doc,
        [
            "Use Sponsor Intake to create a new sponsor manually.",
            "Public self-registration creates pending sponsor records after email verification.",
            "Sponsor details include contact information, status, interest status, drop-off status, and notes.",
            "Sponsor drawer sections are collapsible so staff can focus on communication history and current gift commitments.",
        ],
    )
    add_heading(doc, "Sponsor Directory and Interaction Log", level=2)
    add_body(
        doc,
        "Click a sponsor row to open the sponsor drawer. The interaction log is used to record calls, emails, notes, and follow-up dates. The drawer also shows last contacted information so staff can quickly decide who needs attention.",
    )
    add_heading(doc, "Send a Sponsor Email", level=2)
    add_steps(
        doc,
        [
            "Open Sponsors, then Directory.",
            "Select the sponsor row.",
            "Open the Communication area in the sponsor drawer.",
            "Choose a campaign communication template designed for sponsors.",
            "Preview the message and verify the sponsor has an email address.",
            "Send the email. The send is logged and the sponsor interaction history is updated.",
        ],
    )
    add_heading(doc, "Sponsor Reports", level=2)
    add_body(
        doc,
        "Sponsor Reports summarize sponsor counts, contactable sponsors, public registrations, self-registered sponsors, active sponsorships, sponsored gifts, drop-off status, follow-up queue, and pending public registrations. These reports can be exported.",
    )

    add_heading(doc, "Gifts", level=1)
    add_body(
        doc,
        "The Gifts section manages searching, reserving, receiving, wrapping, tagging, pickup, distribution, and gift pool donations. Gift workflow is tied to wishlist items when the gift is tracked for a specific person.",
    )
    add_screenshot(doc, "gift-search.png", "Gift Search supports natural-language searching by age, gender, gift type, and other filters.")
    add_heading(doc, "Gift Search", level=2)
    add_body(
        doc,
        "Gift Search helps staff or eligible sponsors find gifts to reserve or commit to. The search supports natural-language prompts such as 'girls age 8 to 10 who need coats' and returns matching wishlist items. From there, a user with the right access can commit a sponsor to a gift.",
    )
    add_heading(doc, "Gift Operations", level=2)
    add_bullets(
        doc,
        [
            "Receive gifts when sponsors bring them in.",
            "Move gifts through wrapping, ready, picked up, distributed, exception, and reprint workflows.",
            "Print gift tags for selected gifts or batches.",
            "Print blank/manual tags when staff need a quick tag for a gift that may be filled out by hand.",
        ],
    )
    add_heading(doc, "Gift Pool", level=2)
    add_body(
        doc,
        "The Gift Pool is for donated gifts that are not initially tied to a specific wishlist item, such as a group donation of coats. Staff can inventory those gifts and later match them to people who need them.",
    )
    add_screenshot(doc, "gift-status.png", "Gift Status gives a visual one-page workflow view of recipients, gifts, and status progression.")
    add_heading(doc, "Gift Status Report", level=2)
    add_body(
        doc,
        "Gift Status is the visual report for the gift workflow. It lists recipients, their wishlist gifts, and each status in the workflow so staff can see what is sponsored, received, wrapped, ready, picked up, and distributed at a glance. The page polls while visible so scan updates appear without a manual refresh.",
    )
    add_workflow(
        doc,
        "Update a Gift from Gift Status",
        [
            "Open Gifts, then Gift Status.",
            "Filter the report if needed.",
            "Click a recipient or gift row to open the drawer.",
            "Change status, commit a gift, print a gift tag, or update workflow details.",
            "Save and return to the report.",
        ],
    )
    add_heading(doc, "QR Scan Page", level=2)
    add_body(
        doc,
        "Gift tags include a QR code. Scanning the QR code opens a mobile-friendly public scan page for that gift label. Staff can use the page in a pickup or distribution setting to mark a gift as picked up or distributed, confirm the update, and move to the next scan.",
    )
    add_heading(doc, "Gift Tag Builder", level=2)
    add_body(
        doc,
        "Gift Tag Builder lets campaign managers design the tag template for a campaign. The default tag is 3 inches wide by 2 inches tall, includes the Blessing Tree logo, and must include the QR code. Templates support merge fields such as recipient label, family or organization, age, gender, campaign purpose, and optional gift details.",
    )
    add_bullets(
        doc,
        [
            "QR code is required and cannot be omitted from a valid tracked tag.",
            "Images can be uploaded and placed on the tag.",
            "The default template avoids gift description so the physical gift tag does not reveal gift contents.",
            "Batch printing can request a quantity of tracked or blank tags.",
            "Blank/manual tags create unassigned QR labels for quick hand-written use.",
        ],
    )

    add_heading(doc, "Ask Blessing Tree", level=1)
    add_body(
        doc,
        "Ask Blessing Tree is a conversational help and reporting screen. It can answer app-help questions such as where to add a sponsor, and it can run supported campaign data questions such as how many people still need gifts.",
    )
    add_screenshot(doc, "ask-blessing-tree.png", "Ask Blessing Tree supports help, navigation, and report-style prompts.")
    add_table(
        doc,
        ["Ask Type", "Examples"],
        [
            ["Where do I go?", "Where is Gift Status? How do I add a sponsor? How do I manage organization types?"],
            ["Campaign reporting", "How many children still need sponsors? Show unsponsored gifts. Which sponsors have gifts not received?"],
            ["Operational follow-up", "Show ready gifts not distributed. Show pending public sponsor registrations. Show overdue sponsor gifts."],
            ["Recent work", "Pick up where I left off."],
        ],
        widths=[1.7, 4.55],
    )
    add_note(
        doc,
        "Confidence and review",
        "Low-confidence or negatively rated Ask responses can be reviewed by admins in Admin > Ask Review. That lets the help catalog improve over time.",
    )

    add_heading(doc, "Reports and Exports", level=1)
    add_body(
        doc,
        "Reports appear under People, Sponsors, and Gifts. Report screens are designed to be exported to Excel and PDF. Excel is useful for filtering and analysis; PDF is useful for sharing a clean snapshot.",
    )
    add_table(
        doc,
        ["Report", "What it shows"],
        [
            ["People Reports", "Group mix, people mix, wishlist readiness, people still needing gifts, and pickup coordination."],
            ["Sponsor Reports", "Sponsor counts, drop-off status, follow-up queue, pending public registrations, and sponsored gift totals."],
            ["Gift Status", "Recipient-by-recipient visual gift workflow status with actions from the drawer."],
            ["Dashboard Widgets", "Popular gifts by gender, recipients sponsored by sponsor, unsponsored gifts, counts, and recent Ask prompts."],
            ["Ask Blessing Tree", "Natural-language report answers and report links for supported catalog questions."],
        ],
        widths=[1.7, 4.55],
    )

    add_heading(doc, "Admin", level=1)
    add_body(
        doc,
        "Admin is for app administrators. Most campaign users will not see these screens unless they have app admin permissions.",
    )
    add_screenshot(doc, "admin-user-management.png", "Admin User Management controls users, invitations, active status, and campaign screen access.")
    add_heading(doc, "User Management", level=2)
    add_bullets(
        doc,
        [
            "Invite users and resend invitations.",
            "Activate, deactivate, and delete deactivated users.",
            "Assign app-level role and campaign-level access.",
            "Use the campaign access toggle grid to decide which screens the user sees.",
            "Active campaigns appear first and expanded; inactive campaigns are collapsed.",
        ],
    )
    add_heading(doc, "Campaign Operations", level=2)
    add_body(
        doc,
        "Campaign Operations is the admin rule builder for global milestone definitions and readiness rules. Use it to define which milestones exist, whether they are blockers, where readiness warnings appear, and what message staff see.",
    )
    add_heading(doc, "Organization Types", level=2)
    add_body(
        doc,
        "Organization Types controls the dropdown used by People Intake. Each type has a code, label, active flag, and People Served value. People Served can be Children, Adults, or Families and controls the intake language and flow.",
    )
    add_heading(doc, "LLM Configuration", level=2)
    add_body(
        doc,
        "LLM Configuration controls the model used by Campaign Studio AI and Ask Blessing Tree when LLM features are enabled. Admins can select provider, endpoint, model, and test connectivity.",
    )
    add_heading(doc, "Health Check and App Capabilities", level=2)
    add_body(
        doc,
        "Health Check shows database, Celery, and LLM health. App Capabilities lets admins enable or disable major app surfaces such as People, Sponsors, Reports, Donations, and campaign AI.",
    )

    add_heading(doc, "Account Profile and Settings", level=1)
    add_body(
        doc,
        "Use the menu under your name in the top bar to open Profile or Settings. Profile stores your display information and supports password changes for local accounts. The password fields include visibility controls because the customer requested that behavior. Settings holds account preferences that are separate from campaign setup.",
    )
    add_heading(doc, "Public Sponsor and Public Scan Pages", level=1)
    add_body(
        doc,
        "Public pages are designed for people outside the staff app. Public sponsor registration collects sponsor details first, sends a verification email, and only shows gift selection after verification. Public gift scan is designed for phone use during pickup or distribution and shows only the recipient/gift information and available action buttons.",
    )

    doc.add_page_break()
    add_field_reference_section(doc)
    doc.add_page_break()

    add_heading(doc, "Common Workflows", level=1)
    add_workflow(
        doc,
        "Build a New Campaign",
        [
            "Create or clone the campaign from Campaigns.",
            "Open Campaign Studio and complete Settings.",
            "Add team members and app access in Team.",
            "Create communications templates.",
            "Add milestones and scheduled communications.",
            "Set Gift Rules.",
            "Review Readiness and clear blockers.",
            "Create or review sponsor flyer and gift tag templates.",
        ],
    )
    add_workflow(
        doc,
        "Enter a Family and Wishlist",
        [
            "Open People Intake.",
            "Create the family or household.",
            "Add the first child.",
            "Open the child and add wishlist gifts.",
            "Return to the family drawer and add another child if needed.",
            "Add pickup contacts before gifts are ready for pickup.",
        ],
    )
    add_workflow(
        doc,
        "Receive and Distribute Gifts",
        [
            "Use Gift Search or Sponsor Directory to see committed gifts.",
            "When the gift arrives, mark it received in Gift Operations or Gift Status.",
            "Wrap the gift and print the gift tag.",
            "Mark the gift ready.",
            "Scan the QR tag during pickup or distribution.",
            "Use the mobile action to mark picked up or distributed.",
        ],
    )
    add_workflow(
        doc,
        "Send Sponsor Reminders",
        [
            "Create a sponsor reminder template in Campaign Studio Communications.",
            "Include gift merge fields such as awaiting turn-in list or status list.",
            "Send a test email to yourself.",
            "Open the sponsor from Sponsor Directory.",
            "Select the template from the Communication area and send it to that sponsor.",
            "Review the send record and interaction log.",
        ],
    )

    add_heading(doc, "Glossary", level=1)
    add_table(
        doc,
        ["Term", "Meaning"],
        [
            ["Campaign", "One seasonal or purpose-based giving effort."],
            ["Campaign Purpose", "The campaign theme or purpose that can influence labels, flyers, and gift tag styling."],
            ["Team", "Operational group of campaign workers, such as managers, volunteers, wrapping team, or pickup team."],
            ["App Access", "Permission to see and use specific screens or actions."],
            ["Milestone", "A dated campaign event used for planning, scheduling, and readiness."],
            ["Readiness Rule", "A check that produces a blocker, warning, or note when campaign setup is missing."],
            ["Household", "A family group record, usually with children and pickup contacts."],
            ["Organization", "A partner or program group such as nursing home, foster care, mental health clients, homebound, or other provider."],
            ["People Served", "Admin setting on organization type that tells the app whether the organization flow uses children, adults, or families."],
            ["Wishlist", "The gifts requested for a child or adult."],
            ["Gift Pool", "Unassigned donated gifts that can later be matched to a wishlist."],
            ["Gift Status", "Visual workflow report showing each gift's progress."],
            ["Distributed", "Final gift workflow status used when the gift has been given to the intended person or group."],
        ],
        widths=[1.55, 4.7],
    )

    doc.add_section(WD_SECTION.CONTINUOUS)
    section = doc.sections[-1]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Blessing Tree User Guide").font.size = Pt(9)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_document()
